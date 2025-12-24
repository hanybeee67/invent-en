import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()
    
    # Title
    title = doc.add_heading('Everest Inventory Management System - User Manual', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('एभरेस्ट इन्भेन्टरी व्यवस्थापन प्रणाली - प्रयोगकर्ता पुस्तिका', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')

    # Introduction
    doc.add_heading('1. Introduction (परिचय)', level=1)
    p = doc.add_paragraph()
    p.add_run('KO: ').bold = True
    p.add_run('에베레스트 레스토랑의 재고를 체계적으로 관리하기 위한 시스템입니다.')
    p = doc.add_paragraph()
    p.add_run('NE: ').bold = True
    p.add_run('एभरेस्ट रेस्टुरेन्टको स्टक (सामान) व्यवस्थित रूपमा व्यवस्थापन गर्नको लागि यो एउटा प्रणाली हो।')

    # Access & Login
    doc.add_heading('2. Access & Login (पहुँच र लगइन)', level=1)
    doc.add_paragraph('KO: 첫 화면에서 아무 곳이나 터치하여 진입합니다.', style='List Bullet')
    doc.add_paragraph('NE: पहिलो स्क्रिनमा जहाँसुकै थिचेर भित्र जानुहोस्।', style='List Bullet')
    
    p = doc.add_paragraph('Manager Password (प्रबन्धक पासवर्ड): ', style='List Bullet')
    p.add_run('1234').bold = True
    
    doc.add_paragraph('KO: 분석, 보고서, 데이터 관리 탭은 로그인이 필요합니다.', style='List Bullet 2')
    doc.add_paragraph('NE: विश्लेषण (Analysis), रिपोर्ट (Report), र डाटा व्यवस्थापन (Data Management) ट्याबहरूको लागि लगइन आवश्यक छ।', style='List Bullet 2')

    # Main Features
    doc.add_heading('3. Main Features (मुख्य विशेषताहरू)', level=1)

    # Tab 1
    doc.add_heading('Tab 1: Register / Edit (재고 등록 및 수정 / दर्ता र सम्पादन)', level=2)
    doc.add_paragraph('KO: 품목을 새로 등록하거나 최소 필요 수량을 설정합니다.')
    doc.add_paragraph('NE: नयाँ सामान दर्ता गर्न वा न्यूनतम आवश्यक मात्रा सेट गर्न यहाँ जानुहोस्।')

    # Tab 2
    doc.add_heading('Tab 2: View / Print (재고 조회 및 출력 / स्टक हेर्ने र प्रिन्ट गर्ने)', level=2)
    doc.add_paragraph('KO: 현재 재고 현황을 지점별, 카테고리별로 확인하고 인쇄용 파일로 다운로드합니다.')
    doc.add_paragraph('NE: शाखा वा श्रेणी अनुसार हालको स्टक विवरण हेर्नुहोस् र प्रिन्ट गर्नको लागि फाइल डाउनलोड गर्नुहोस्।')

    # Tab 3
    doc.add_heading('Tab 3: IN / OUT Log (입출고 기록 / भित्र र बाहिरको रेकर्ड)', level=2)
    doc.add_paragraph('KO: 배송된 물품(IN)이나 사용한 물품(OUT)을 기록하면 재고 수량이 자동으로 업데이트됩니다.')
    doc.add_paragraph('NE: आएको सामान (IN) वा प्रयोग भएको सामान (OUT) रेकर्ड गर्नुहोस्। यसले स्टकको मात्रा आफैं मिलाउनेछ।')

    # Tab 4
    doc.add_heading('Tab 4: Usage Analysis (사용량 분석 / प्रयोगको विश्लेषण)', level=2)
    doc.add_paragraph('KO: 어느 지점에서 어떤 품목이 많이 사용되었는지 분석합니다.')
    doc.add_paragraph('NE: कुन शाखामा कुन सामान धेरै प्रयोग भइरहेको छ भनेर यहाँ हेर्न सकिन्छ।')

    # Tab 5
    doc.add_heading('Tab 5: Monthly Report (월간 보고서 / मासिक रिपोर्ट)', level=2)
    doc.add_paragraph('KO: 월별 재고 및 사용량 내역을 엑셀(Excel)이나 PDF로 저장합니다.')
    doc.add_paragraph('NE: हरेक महिनाको स्टक र प्रयोगको विवरण एक्सेल वा PDF मा डाउनलोड गर्नुहोस्।')

    # Tab 6
    doc.add_heading('Tab 6: Data Management (데이터 관리 / 데이터 व्यवस्थापन)', level=2)
    doc.add_paragraph('KO: 엑셀 파일을 업로드하여 품목 리스트를 한꺼번에 업데이트합니다.')
    doc.add_paragraph('NE: एक्सेल फाइल अपलोड गरेर सबै सामानको सूची एकैपटक अपडेट गर्नुहोस्।')

    # Special Features
    doc.add_heading('4. Special Features (विशेष विशेषताहरू)', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Low Stock Alert (재고 부족 알림 / कम स्टकको सूचना):').bold = True
    doc.add_paragraph('KO: 재고가 설정된 최소 수량 이하로 떨어지면 빨간색 경고창이 뜹니다.', style='List Bullet')
    doc.add_paragraph('NE: यदि सामानको मात्रा न्यूनतम सेट गरिएको भन्दा कम भयो भने, रातो रङ्गको चेतावनी देखिनेछ।', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Persistence (데이터 영구 저장 / डाटा सुरक्षित):').bold = True
    doc.add_paragraph('KO: 데이터는 서버에 자동으로 저장되므로 앱을 꺼도 사라지지 않습니다.', style='List Bullet')
    doc.add_paragraph('NE: डाटा आफैं सुरक्षित हुनेछ, त्यसैले एप बन्द गरे पनि मेटिने छैन।', style='List Bullet')

    # Save
    base_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(base_dir, 'Everest_Manual.docx')
    doc.save(file_path)
    print(f"Manual saved to: {file_path}")

if __name__ == "__main__":
    create_manual()
